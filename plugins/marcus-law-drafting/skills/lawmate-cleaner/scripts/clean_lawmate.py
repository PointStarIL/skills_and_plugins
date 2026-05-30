#!/usr/bin/env python3
"""
Lawmate Cleaner - עיבוד טיוטות משפטיות שהופקו ממערכת law-mate.

הסקריפט מנקה את הטקסט (LawMate stamps, file refs, em-dashes, LRM/RLM marks,
date normalisation, citation deduplication) ובונה DOCX מעוצב דרך המנוע המשותף
docx-hebrew-engine, שמחזיק את התבנית והסגנונות בשם: Normal, List Paragraph
(decimal auto-numbering), Heading 2, hebrew1 numbering לסעדים.

מבנה הפלט:
- פסקת גוף: List Paragraph (מספור 1, 2, 3 רץ דרך style inheritance)
- אזכור נספח ("מצורף ומסומן כנספח N"): List Paragraph + numId=0 (override),
  עם קו תחתון
- סעד ("א.", "ב.", "ג." בתחילת פסקה): List Paragraph + hebrew1 numbering,
  הקידומת מוסרת
- כותרת-משנה (centered + bold בקלט): Heading 2
"""

import argparse
import os
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Shared engine: the docx-hebrew-engine skill is the single source of truth
# for styles/template/RTL — all DOCX generation routes through it.
sys.path.insert(0, str(
    Path(__file__).resolve().parents[2] / "docx-hebrew-engine" / "scripts"))
import docx_hebrew_engine as hd
from docx_hebrew_engine import (
    BOLD_OPEN, BOLD_CLOSE, STYLE_BODY, STYLE_HEADING,
    EXHIBIT_NUMID, REMEDY_NUMID,
)

# EMU thresholds for heading detection in lawmate input
MAIN_HEADING_EMU = 177800   # 14pt
SUB_HEADING_EMU = 152400    # 12pt

# Detection patterns
EXHIBIT_RE = re.compile(r'מצורף\s+ומסומן\s+כנספח\s+\d+')
REMEDY_PREFIX_RE = re.compile(r'^([א-ת])\.\s+(.+)$', re.DOTALL)
# Splits an exhibit sentence into (description, number):
# "העתק X מיום ... מצורף ומסומן כנספח 3." -> desc="העתק X מיום ...", num=3
EXHIBIT_PARSE_RE = re.compile(r'^(.*?)\s*מצורף\s+ומסומן\s+כנספח\s+(\d+)')

EXHIBIT_LIST_HEADING = "רשימת הנספחים"


def _bold_parties(parties: str) -> str:
    parties = parties.strip().strip(',').strip()
    if not re.search(r'[א-ת]', parties):
        return parties
    return f'{BOLD_OPEN}{parties}{BOLD_CLOSE}'


def _format_full_citation(m):
    court = m.group(1).strip()
    case_id = m.group(2).strip()
    parties = m.group(3).strip()
    date = m.group(4).strip()
    if not re.search(r'[א-ת]', parties.replace('נ', '')):
        return ''
    return f'({court}, {case_id}, {_bold_parties(parties)}, {date})'


def _format_lawmate_citation(m):
    inner = m.group(1).strip()
    date = m.group(2).strip()
    parts = inner.split(None, 2)
    if len(parts) == 3:
        prefix, case_id, parties = parts
        if re.search(r'[א-ת]', parties.replace('נ', '')):
            return f'{prefix} {case_id} {_bold_parties(parties)} ({date})'
    return f'{inner} ({date})'


def _format_paren_lawmate_citation(m):
    """Format: (PREFIX CASE_ID PARTIES (DATE LawMate)) -> (PREFIX CASE_ID **PARTIES** (DATE))"""
    prefix = m.group(1).strip()
    case_id = m.group(2).strip()
    parties = m.group(3).strip().rstrip(',').strip()
    date = m.group(4).strip()
    if not re.search(r'[א-ת]', parties.replace('נ', '')):
        return ''
    return f'({prefix} {case_id} {_bold_parties(parties)} ({date}))'


def clean_text(t: str) -> str:
    if not t:
        return t

    # -1. Strip directional marks (LRM/RLM) that fragment regex matching
    t = t.replace('‎', '').replace('‏', '')

    # 0. Normalise dates DD/MM/YYYY -> DD.MM.YYYY. Case numbers (e.g.
    #    64925-09-25) and law years (e.g. התשנ"ה-1995) use hyphens, so they
    #    are never touched by this slash-only pattern.
    t = re.sub(r'\b(\d{1,2})/(\d{1,2})/(\d{4})\b', r'\1.\2.\3', t)

    # 1. Remove trailing appendix/page references ] (12) -> ]
    t = re.sub(r'\]\s*\(\d{1,4}\)', ']', t)

    # 2. Replace em/en dashes with hyphens (AI tell)
    t = t.replace('—', '-').replace('–', '-')

    # 3. Full-format citations: [court, case_id, parties, date]
    t = re.sub(
        r'\[([^,\]]+),\s*([^,\]]+),\s*([^\]]*?נ[\'.]\s*[^,\]]*?),\s*([\d./]+)\]',
        _format_full_citation, t,
    )

    # 4. Case-law citations with LawMate stamp
    case_prefixes = ['עב"ל', 'ע"א', 'בג"ץ', 'בג"צ', 'ב"ל', 'ע"ע', 'תיק', 'בר"ע']
    for prefix in case_prefixes:
        # 4a. Bracket format: [PREFIX ... (LawMate DATE)]
        t = re.sub(
            r'\[(' + re.escape(prefix) + r'[^\]]*?)\(LawMate\s+([^)]+)\)\]',
            _format_lawmate_citation, t)
        # 4b. Bracket format: [PREFIX ... (DATE LawMate)]
        t = re.sub(
            r'\[(' + re.escape(prefix) + r'[^\]]*?)\(([\d./]+)\s+LawMate\)\]',
            _format_lawmate_citation, t)
        # 4c. Paren format: (PREFIX CASE_ID PARTIES (DATE LawMate))
        t = re.sub(
            r'\(\s*(' + re.escape(prefix) + r')\s+([^\s()]+)\s+([^()]+?)\s*\(\s*([\d./]+)\s+LawMate\s*\)\s*\)',
            _format_paren_lawmate_citation, t)
        # 4d. Paren format: (PREFIX CASE_ID PARTIES (LawMate DATE))
        t = re.sub(
            r'\(\s*(' + re.escape(prefix) + r')\s+([^\s()]+)\s+([^()]+?)\s*\(\s*LawMate\s+([\d./]+)\s*\)\s*\)',
            _format_paren_lawmate_citation, t)

    # 5. File references
    t = re.sub(r'\s*\[תיק-\d+[^\]]*\]', '', t)
    t = re.sub(r'\s*\[[^\]]*\.(pdf|docx)[^\]]*\]', '', t)
    t = re.sub(r"\s*\[[^\]]*?עמ'[^\]]*\]", '', t)

    # 6. General law-citation brackets
    t = re.sub(r'\s*\[חוק[^\]]+\]', '', t)

    # 7. Bare (N) appendix references at end of sentence
    t = re.sub(r'\s+\(\d{1,4}\)(?=\s*[.,;:]|\s*$)', '', t)

    # 8. Whitespace cleanup (don't touch \x01/\x02)
    t = re.sub(r'[ \t]+', ' ', t)
    t = re.sub(r'\s+([.,;:])', r'\1', t)
    t = re.sub(r'\(\s+', '(', t)
    t = re.sub(r'\s+\)', ')', t)
    return t.strip()


def classify_paragraph(p):
    """Classify a paragraph into one of: sub_heading, exhibit_ref, remedy, body.

    Returns (kind, cleaned_text) or None for empty paragraphs.
    """
    raw = p.text.strip()
    if not raw:
        return None

    align = p.alignment
    bold = False
    size = None
    if p.runs:
        r = p.runs[0]
        bold = bool(r.font.bold)
        size = r.font.size
    style_name = p.style.name if p.style else ""

    is_centered = align == WD_ALIGN_PARAGRAPH.CENTER
    is_heading = (
        (is_centered and bold and (size is None or size >= SUB_HEADING_EMU))
        or style_name in ("Heading 1", "Heading 2")
    )
    if is_heading:
        return ("sub_heading", clean_text(raw))

    cleaned = clean_text(raw)
    if not cleaned:
        return None

    if EXHIBIT_RE.search(cleaned):
        return ("exhibit_ref", cleaned)

    m = REMEDY_PREFIX_RE.match(cleaned)
    if m:
        return ("remedy", m.group(2).strip())

    return ("body", cleaned)


def extract_items(doc):
    """Walk the lawmate doc, skip the title-block envelope, classify each
    substantive paragraph. The body begins at the first heading."""
    items = []
    in_body = False
    for p in doc.paragraphs:
        c = classify_paragraph(p)
        if c is None:
            continue
        kind, text = c
        if not text:
            continue
        if not in_body:
            if kind == "sub_heading":
                in_body = True
            else:
                continue
        items.append((kind, text))
    return items


_CITE_PREFIXES = '(?:עב"ל|ע"א|בג"ץ|בג"צ|ב"ל|ע"ע|תיק|בר"ע)'
_CITE_RE = re.compile(
    r'\(\s*(' + _CITE_PREFIXES + r')\s+'
    r'([\S]+?)\s+'
    + re.escape(BOLD_OPEN) + r'([^' + re.escape(BOLD_CLOSE) + r']+)' + re.escape(BOLD_CLOSE)
    + r'\s*(?:\(\s*([\d./]+)\s*\)|,\s*([\d./]+))\s*\)'
)


def _tidy_after_removal(t: str) -> str:
    """Clean up punctuation/whitespace left behind when a citation is deleted."""
    t = re.sub(r'\(\s*\)', '', t)        # empty parens
    t = re.sub(r'\s{2,}', ' ', t)        # collapse double spaces
    t = re.sub(r'\s+([.,;:])', r'\1', t)  # space before punctuation
    t = re.sub(r';\s*;', ';', t)          # orphaned ;;
    t = re.sub(r';\s*([.)])', r'\1', t)   # "; ." -> "."
    t = re.sub(r'\(\s*;', '(', t)         # "( ;"
    return t.strip()


def dedup_citations(items):
    """De-duplicate repeated case-law citations.

    A case is cited in full the first time it appears. Every later
    parenthetical citation of the same case is removed entirely — the case
    is already established, and the prose itself carries the name where it
    matters ("הלכת עותמאן", "בעניין דורון כושאווי"). Leaving a bare
    "עניין X" in place of the dropped citation reads as a dangling fragment,
    so we delete the whole parenthetical and tidy the surrounding punctuation.
    """
    seen = set()

    def process(text):
        removed_any = False
        result = []
        pos = 0
        for m in _CITE_RE.finditer(text):
            result.append(text[pos:m.start()])
            case_id = m.group(2)
            if case_id not in seen:
                seen.add(case_id)
                result.append(m.group(0))   # first mention: keep full citation
            else:
                removed_any = True           # repeat: drop entirely
            pos = m.end()
        result.append(text[pos:])
        out = ''.join(result)
        return _tidy_after_removal(out) if removed_any else out

    out = []
    for kind, text in items:
        if kind in ('body', 'exhibit_ref', 'remedy'):
            text = process(text)
        out.append((kind, text))
    return out


# ----- DOCX building (delegates to the shared docx_hebrew_engine) --------


def _parse_exhibit(text):
    """From an exhibit sentence, return (number:int, description:str).

    "העתק פסק הדין ... מצורף ומסומן כנספח 2." -> (2, "העתק פסק הדין ...")
    Returns None if the sentence doesn't match the expected shape.
    """
    plain = text.replace(BOLD_OPEN, '').replace(BOLD_CLOSE, '')
    m = EXHIBIT_PARSE_RE.match(plain)
    if not m:
        return None
    desc = m.group(1).strip()
    return (int(m.group(2)), desc)


def _add_exhibit_list(doc, exhibits):
    """Append a consolidated exhibit index at the end of the document:
    a Heading 2 'רשימת הנספחים' followed by one line per exhibit, ordered
    by number, each reading 'נספח N - <description>' with 'נספח N' bold."""
    if not exhibits:
        return
    heading = doc.add_paragraph(style=STYLE_HEADING)
    hd.add_run(heading, EXHIBIT_LIST_HEADING)
    for num, desc in sorted(exhibits, key=lambda e: e[0]):
        p = doc.add_paragraph(style=STYLE_BODY)
        hd.set_numbering(p, EXHIBIT_NUMID)   # override → no auto list marker
        hd.add_run(p, f'נספח {num} - ', bold=True)
        hd.add_run(p, desc)


def build_docx(items, output_path):
    doc = hd.open_document()

    exhibits = []
    for kind, text in items:
        if kind == "sub_heading":
            hd.add_heading(doc, text)
            continue

        if kind == "exhibit_ref":
            hd.add_exhibit_ref(doc, text)
            parsed = _parse_exhibit(text)
            if parsed:
                exhibits.append(parsed)
            continue

        if kind == "remedy":
            hd.add_hebrew_item(doc, text)
            continue

        # body
        hd.add_body(doc, text)

    # Consolidated exhibit index at the very end
    _add_exhibit_list(doc, exhibits)

    hd.save(doc, output_path)


def main():
    parser = argparse.ArgumentParser(description="Clean a law-mate DOCX draft.")
    parser.add_argument("input", help="Path to input DOCX from law-mate")
    parser.add_argument("output", nargs="?",
                        help="Path for cleaned DOCX")
    args = parser.parse_args()

    if not os.path.exists(args.input):
        print(f"Error: input file not found: {args.input}", file=sys.stderr)
        sys.exit(1)

    output = args.output
    if output is None:
        base, _ = os.path.splitext(args.input)
        output = f"{base}_מסודר.docx"

    print(f"Reading: {args.input}")
    src = Document(args.input)
    raw_para_count = len(src.paragraphs)

    items = extract_items(src)
    items = dedup_citations(items)

    sub_count = sum(1 for k, _ in items if k == "sub_heading")
    body_count = sum(1 for k, _ in items if k == "body")
    exhibit_count = sum(1 for k, _ in items if k == "exhibit_ref")
    remedy_count = sum(1 for k, _ in items if k == "remedy")

    full_raw_text = "\n".join(p.text for p in src.paragraphs)
    full_clean_text = "\n".join(t for _, t in items)
    raw_brackets = len(re.findall(r'\[[^\]]+\]', full_raw_text))
    clean_brackets = len(re.findall(r'\[[^\]]+\]', full_clean_text))
    lawmate_raw = len(re.findall(r'(?i)lawmate', full_raw_text))
    lawmate_clean = len(re.findall(r'(?i)lawmate', full_clean_text))
    raw_emdash = full_raw_text.count('—') + full_raw_text.count('–')
    clean_emdash = full_clean_text.count('—') + full_clean_text.count('–')
    bold_count = full_clean_text.count(BOLD_OPEN)

    build_docx(items, output)

    print(f"\n  Wrote: {output}")
    print(f"  Paragraphs in source: {raw_para_count}")
    print(f"  Items in output: {len(items)}")
    print(f"  Sub-headings (Heading 2): {sub_count}")
    print(f"  Body paragraphs (numbered 1,2,3): {body_count}")
    print(f"  Exhibit references (underlined): {exhibit_count}")
    print(f"  Remedies (א, ב, ג, ד): {remedy_count}")
    print(f"  Bracket refs: {raw_brackets} -> {clean_brackets}")
    print(f"  LawMate mentions: {lawmate_raw} -> {lawmate_clean}")
    print(f"  Em/en dashes: {raw_emdash} -> {clean_emdash}")
    print(f"  Bolded party citations: {bold_count}")


if __name__ == "__main__":
    main()
