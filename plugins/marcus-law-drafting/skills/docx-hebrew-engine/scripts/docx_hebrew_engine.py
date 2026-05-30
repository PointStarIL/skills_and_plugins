#!/usr/bin/env python3
"""
docx_hebrew_engine - מנוע משותף לבניית מסמכי Word משפטיים בעברית (python-docx).

מקור-אמת יחיד לעיצוב: כל הסגנונות, הגדרות המספור, השוליים וה-RTL יושבים
ב-`template.docx` שלצד הקובץ הזה. כל מסמך שנבנה דרך המנוע פותח את התבנית
ומחיל סגנונות בשם - כך כל המסמכים (lawmate-cleaner, legal-docx, וכל סקריפט
עתידי) יוצאים עם עיצוב זהה.

עקרון ה-RTL הקריטי: סגנון "List Paragraph" בתבנית נושא <w:rtl w:val="0"/>,
שמכבה RTL. לכן כל run מקבל <w:rtl/> מפורש שדורס זאת - אחרת Word מתייחס
לעברית כ-LTR ומציג אותה בפונט ה-ascii (Times New Roman) במקום David.
שמות מודגשים מקבלים גם <w:bCs/> (bold complex-script - עברית מתבלטת רק כך).
אין לקבוע rFonts/sz ברמת ה-run: המיפוי ascii=Times New Roman / cs=David
והגדלים מגיעים מהסגנון.

שימוש לדוגמה (בנייה מאפס):

    import sys
    from pathlib import Path
    sys.path.insert(0, str(Path(__file__).resolve().parents[N] / "shared"))
    import docx_hebrew_engine as hd

    doc = hd.open_document()
    hd.add_heading(doc, "רקע עובדתי ותיאור הצדדים")
    hd.add_body(doc, "פסקה ראשונה...")      # ממוספר 1
    hd.add_body(doc, "פסקה שנייה...")       # ממוספר 2
    hd.add_hebrew_item(doc, "סעד מבוקש ראשון")  # א.
    hd.add_exhibit_ref(doc, "העתק X מצורף ומסומן כנספח 1.")  # קו תחתון
    hd.save(doc, "out.docx")
"""

import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# Inline bold markers: wrap a span in BOLD_OPEN..BOLD_CLOSE and it renders bold.
BOLD_OPEN = '\x01'
BOLD_CLOSE = '\x02'

TEMPLATE_PATH = Path(__file__).resolve().parent.parent / "references" / "template.docx"

# Style names (resolved by python-docx via name → styleId lookup in the template)
STYLE_BODY = "List Paragraph"     # auto-numbered 1, 2, 3 (numId=14 via style)
STYLE_HEADING = "Heading 2"

# numId values predefined in template.docx's numbering.xml
EXHIBIT_NUMID = 0    # override → "no list" marker (used for underlined exhibit refs)
REMEDY_NUMID = 43    # → abstractNumId=7 (hebrew1: א., ב., ג., ד.)

# Directional marks that fragment regex/search when ingesting source text.
_LRM = '‎'
_RLM = '‏'


def sanitize_source_text(t: str) -> str:
    """Clean ingested source text before placing it in a document:
    strip LRM/RLM directional marks and replace em/en dashes (AI tells)
    with a plain hyphen. Safe to call on any user/AI-supplied string."""
    if not t:
        return t
    t = t.replace(_LRM, '').replace(_RLM, '')
    t = t.replace('—', '-').replace('–', '-')
    return t


def open_document():
    """Open the canonical template and strip its placeholder paragraph,
    returning a Document ready for content. Raises if the template is missing."""
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Template not found: {TEMPLATE_PATH}")
    doc = Document(str(TEMPLATE_PATH))
    if doc.paragraphs:
        placeholder = doc.paragraphs[0]
        placeholder._element.getparent().remove(placeholder._element)
    return doc


def set_numbering(p, num_id, ilvl=0):
    """Apply an explicit numPr to a paragraph (overrides the style's numbering)."""
    pPr = p._p.get_or_add_pPr()
    numPr = OxmlElement('w:numPr')
    ilvl_el = OxmlElement('w:ilvl')
    ilvl_el.set(qn('w:val'), str(ilvl))
    numPr.append(ilvl_el)
    numId_el = OxmlElement('w:numId')
    numId_el.set(qn('w:val'), str(num_id))
    numPr.append(numId_el)
    pPr.append(numPr)


def apply_paragraph_underline(p):
    """Add a paragraph-level single underline (covers an auto list marker)."""
    pPr = p._p.get_or_add_pPr()
    rPr = pPr.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        pPr.append(rPr)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)


def add_run(p, text, bold=False, underline=False):
    """Add a run with an explicit RTL marker. Font/size come from the style.

    Every run gets <w:rtl/> to override the List Paragraph style's rtl=0;
    bold runs also get <w:bCs/> (Hebrew renders bold only via complex-script).
    rFonts is deliberately left unset so the Normal style's
    ascii=Times New Roman / cs=David mapping applies.
    """
    r = p.add_run(text)
    rPr = r._r.get_or_add_rPr()
    if bold:
        r.bold = True
        rPr.append(OxmlElement('w:bCs'))
    if underline:
        r.underline = True
    rPr.append(OxmlElement('w:rtl'))
    return r


def add_formatted_text(p, text, underline=False):
    """Add text to a paragraph, converting inline BOLD_OPEN..BOLD_CLOSE spans
    to bold runs. underline=True applies a single underline to every run."""
    if not text:
        return
    if BOLD_OPEN not in text:
        add_run(p, text, underline=underline)
        return
    pattern = re.compile(re.escape(BOLD_OPEN) + r'(.*?)' + re.escape(BOLD_CLOSE))
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            add_run(p, text[pos:m.start()], underline=underline)
        add_run(p, m.group(1), bold=True, underline=underline)
        pos = m.end()
    if pos < len(text):
        add_run(p, text[pos:], underline=underline)


# ----- High-level paragraph builders -----------------------------------


def add_heading(doc, text):
    """Add a section heading (Heading 2 style). Bold markers are stripped."""
    clean = text.replace(BOLD_OPEN, '').replace(BOLD_CLOSE, '')
    p = doc.add_paragraph(style=STYLE_HEADING)
    add_run(p, clean)
    return p


def add_body(doc, text):
    """Add a body paragraph (List Paragraph, auto-numbered 1, 2, 3...)."""
    p = doc.add_paragraph(style=STYLE_BODY)
    add_formatted_text(p, text)
    return p


def add_hebrew_item(doc, text):
    """Add a paragraph numbered with Hebrew letters (א., ב., ג., ...)."""
    p = doc.add_paragraph(style=STYLE_BODY)
    set_numbering(p, REMEDY_NUMID)
    add_formatted_text(p, text)
    return p


def add_exhibit_ref(doc, text):
    """Add an underlined exhibit reference (no visible list marker)."""
    p = doc.add_paragraph(style=STYLE_BODY)
    set_numbering(p, EXHIBIT_NUMID)
    apply_paragraph_underline(p)
    add_formatted_text(p, text, underline=True)
    return p


def add_plain(doc, text):
    """Add a non-numbered body paragraph (List Paragraph, marker suppressed)."""
    p = doc.add_paragraph(style=STYLE_BODY)
    set_numbering(p, EXHIBIT_NUMID)   # numId=0 → no marker
    add_formatted_text(p, text)
    return p


def save(doc, output_path):
    doc.save(str(output_path))
