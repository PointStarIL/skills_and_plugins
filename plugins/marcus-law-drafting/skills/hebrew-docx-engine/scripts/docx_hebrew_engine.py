#!/usr/bin/env python3
"""
docx_hebrew_engine - מנוע משותף לבניית מסמכי Word משפטיים בעברית (python-docx).

מקור-אמת יחיד לעיצוב: כל הסגנונות, הגדרות המספור, השוליים וה-RTL יושבים
ב-`template.docx` שלצד הקובץ הזה. כל מסמך שנבנה דרך המנוע פותח את התבנית
ומחיל סגנונות בשם - כך כל המסמכים (clean-lawmate-draft, edit-legal-docx, וכל סקריפט
עתידי) יוצאים עם עיצוב זהה.

עקרון ה-RTL הקריטי: כל run מקבל <w:rtl/> מפורש. בלעדיו Word עלול להתייחס
לעברית כ-LTR ולהציג אותה בפונט ה-ascii במקום ב-David. שמות מודגשים מקבלים
גם <w:bCs/> (bold complex-script - עברית מתבלטת רק כך).

אין לקבוע rFonts או sz ברמת ה-run. סגנון Normal בתבנית מפריד בין השפות:
ascii/hAnsi = Times New Roman בגודל 10 (w:sz=20), ו-cs = David בגודל 12
(w:szCs=24). כל קביעה ברמת ה-run תדרוס את ההפרדה ותשבור את גודל האנגלית.

עקרון היישור בתאי טבלה: אין לקבוע w:jc בפסקה שבתוך תא. ב-OOXML הערכים
left/right הם כינויים של start/end, ו-Word ממפה right ל-end. בפסקה עם
<w:bidi/> הקצה הוא צד שמאל, ולכן jc=right מיישר שמאלה. בלי w:jc התא יורש
מ-Normal את jc=both, וזה היישור הנכון.

שימוש לדוגמה (בנייה מאפס):

    import sys
    from pathlib import Path
    sys.path.insert(0, str(Path(__file__).resolve().parents[N] / "shared"))
    import docx_hebrew_engine as hd

    doc = hd.open_document()
    hd.add_heading(doc, "רקע עובדתי ותיאור הצדדים")   # כותרת 2
    hd.add_heading3(doc, "תת פרק")                    # כותרת 3
    hd.add_body(doc, "פסקה ראשונה...")                # 1.
    hd.add_body(doc, "פסקה שנייה...")                 # 2.
    hd.add_hebrew_item(doc, "סעד מבוקש ראשון")        # א.   (סעיף, ilvl=1)
    hd.add_clause(doc, "רמה שלישית", level=2)         # (1)
    hd.add_clause(doc, "רמה רביעית", level=3)         # (א)
    hd.add_quote(doc, "ציטוט מתוך הפרוטוקול")         # סגנון Quote
    hd.add_exhibit_ref(doc, "העתק X מצורף ומסומן כנספח 1.")  # שורת נספח
    hd.save(doc, "out.docx")
"""

import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm
from docx.enum.table import WD_TABLE_ALIGNMENT


# Inline bold markers: wrap a span in BOLD_OPEN..BOLD_CLOSE and it renders bold.
BOLD_OPEN = '\x01'
BOLD_CLOSE = '\x02'

TEMPLATE_PATH = Path(__file__).resolve().parent.parent / "references" / "template.docx"

# שמות הסגנונות. python-docx מאתר אותם לפי השם ולא לפי styleId, ולכן שינוי
# styleId בתבנית אינו מחייב שינוי כאן.
STYLE_BODY = "סעיף"          # מספור אוטומטי דרך numId=5, ארבע רמות הזחה
STYLE_TITLE = "Title"        # כותרת ראשית של המסמך
STYLE_HEADING = "Heading 2"
STYLE_HEADING3 = "Heading 3"
STYLE_EXHIBIT = "Subtitle"   # כינויו בעברית: "שורת נספח"
STYLE_QUOTE = "Quote"

# הרשימה שסגנון "סעיף" נשען עליה. רמות ההזחה:
#   ilvl=0  ->  1.     (decimal)
#   ilvl=1  ->  א.     (hebrew1)
#   ilvl=2  ->  (1)    (decimal)
#   ilvl=3  ->  (א)    (hebrew2)
CLAUSE_NUMID = 5
MAX_CLAUSE_LEVEL = 3

# Directional marks that fragment regex/search when ingesting source text.
_LRM = '‎'
_RLM = '‏'


def sanitize_source_text(t: str) -> str:
    """Clean ingested source text before placing it in a document:
    strip LRM/RLM directional marks and replace em/en dashes (AI tells)
    with a plain hyphen. Safe to call on any user/AI-supplied string."""
    if not t:
        return t
    t = t.replace(_LRM, '').replace(_RLM, '')
    t = t.replace('—', '-').replace('–', '-')
    return t


def open_document():
    """Open the canonical template and strip its placeholder paragraph,
    returning a Document ready for content. Raises if the template is missing."""
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Template not found: {TEMPLATE_PATH}")
    doc = Document(str(TEMPLATE_PATH))
    if doc.paragraphs:
        placeholder = doc.paragraphs[0]
        placeholder._element.getparent().remove(placeholder._element)
    return doc


def set_numbering(p, num_id, ilvl=0):
    """Apply an explicit numPr to a paragraph (overrides the style's numbering)."""
    pPr = p._p.get_or_add_pPr()
    numPr = OxmlElement('w:numPr')
    ilvl_el = OxmlElement('w:ilvl')
    ilvl_el.set(qn('w:val'), str(ilvl))
    numPr.append(ilvl_el)
    numId_el = OxmlElement('w:numId')
    numId_el.set(qn('w:val'), str(num_id))
    numPr.append(numId_el)
    pPr.append(numPr)


def apply_paragraph_underline(p):
    """Add a paragraph-level single underline (covers an auto list marker)."""
    pPr = p._p.get_or_add_pPr()
    rPr = pPr.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        pPr.append(rPr)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)


def add_run(p, text, bold=False, underline=False):
    """Add a run with an explicit RTL marker. Font/size come from the style.

    כל run מקבל <w:rtl/> מפורש; runs מודגשים מקבלים גם <w:bCs/>, מפני
    שעברית מתבלטת רק דרך complex-script. rFonts ו-sz נשארים ללא קביעה
    בכוונה, כדי שההפרדה שבסגנון Normal תחול: ascii = Times New Roman 10,
    cs = David 12.

    הטקסט עובר sanitize_source_text אוטומטית: הסרת LRM/RLM והחלפת מקף
    ארוך/בינוני במקף רגיל. זו נקודת המעבר היחידה שכל טקסט עובר בה בדרכו
    למסמך, ולכן ההגנה כאן חלה על כל הסקיילים בלי תלות במשמעת הקורא.
    """
    text = sanitize_source_text(text)
    r = p.add_run(text)
    rPr = r._r.get_or_add_rPr()
    if bold:
        r.bold = True
        rPr.append(OxmlElement('w:bCs'))
    if underline:
        r.underline = True
    rPr.append(OxmlElement('w:rtl'))
    return r


def add_formatted_text(p, text, underline=False):
    """Add text to a paragraph, converting inline BOLD_OPEN..BOLD_CLOSE spans
    to bold runs. underline=True applies a single underline to every run."""
    if not text:
        return
    if BOLD_OPEN not in text:
        add_run(p, text, underline=underline)
        return
    pattern = re.compile(re.escape(BOLD_OPEN) + r'(.*?)' + re.escape(BOLD_CLOSE))
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            add_run(p, text[pos:m.start()], underline=underline)
        add_run(p, m.group(1), bold=True, underline=underline)
        pos = m.end()
    if pos < len(text):
        add_run(p, text[pos:], underline=underline)


# ----- High-level paragraph builders -----------------------------------


def add_title(doc, text):
    """כותרת ראשית של המסמך (סגנון Title). סימוני הדגשה מוסרים.

    היישור והמרווחים מגיעים מהסגנון שבתבנית. אין לקבוע כאן w:jc: בפסקת
    bidi הערך right ממופה ל-end, שהוא צד שמאל. ראה את הערת היישור בראש הקובץ.
    """
    clean = text.replace(BOLD_OPEN, '').replace(BOLD_CLOSE, '')
    p = doc.add_paragraph(style=STYLE_TITLE)
    add_run(p, clean)
    return p


def add_heading(doc, text):
    """Add a section heading (Heading 2 style). Bold markers are stripped."""
    clean = text.replace(BOLD_OPEN, '').replace(BOLD_CLOSE, '')
    p = doc.add_paragraph(style=STYLE_HEADING)
    add_run(p, clean)
    return p


def add_heading3(doc, text):
    """כותרת משנה (סגנון Heading 3). סימוני הדגשה מוסרים."""
    clean = text.replace(BOLD_OPEN, '').replace(BOLD_CLOSE, '')
    p = doc.add_paragraph(style=STYLE_HEADING3)
    add_run(p, clean)
    return p


def add_clause(doc, text, level=0):
    """פסקת גוף בסגנון "סעיף", ברמת ההזחה המבוקשת (0 עד 3).

    level=0 -> 1.    level=1 -> א.    level=2 -> (1)    level=3 -> (א)

    רמה 0 מגיעה מהסגנון עצמו ואינה דורשת numPr מפורש. ברמות 1 עד 3 מוחל
    numPr על אותה רשימה (numId=5) עם ilvl אחר, בדיוק כפי שנעשה ב-Word.
    """
    if not 0 <= level <= MAX_CLAUSE_LEVEL:
        raise ValueError(f"level חייב להיות בין 0 ל-{MAX_CLAUSE_LEVEL}, התקבל {level}")
    p = doc.add_paragraph(style=STYLE_BODY)
    if level:
        set_numbering(p, CLAUSE_NUMID, ilvl=level)
    add_formatted_text(p, text)
    return p


def add_body(doc, text):
    """פסקת גוף ממוספרת 1, 2, 3 (סגנון "סעיף", רמה 0)."""
    return add_clause(doc, text, level=0)


def add_hebrew_item(doc, text):
    """פסקה ממוספרת באותיות עבריות א., ב., ג. (סגנון "סעיף", רמה 1)."""
    return add_clause(doc, text, level=1)


def add_exhibit_ref(doc, text):
    """שורת הפניה לנספח (סגנון "שורת נספח"/Subtitle).

    הסגנון עצמו מבטל את המספור ומחיל מודגש וקו תחתון, ולכן אין צורך
    בקו תחתון ידני או בביטול מספור בקוד.
    """
    p = doc.add_paragraph(style=STYLE_EXHIBIT)
    add_formatted_text(p, text)
    return p


def add_quote(doc, text):
    """ציטוט (סגנון Quote): מודגש, מוזח משני הצדדים, מרווח 1.15."""
    p = doc.add_paragraph(style=STYLE_QUOTE)
    add_formatted_text(p, text)
    return p


def add_plain(doc, text):
    """פסקה לא ממוספרת בסגנון הרגיל (Normal)."""
    p = doc.add_paragraph()
    add_formatted_text(p, text)
    return p


# ----- RTL tables ------------------------------------------------------
#
# RTL table principle (mirrors the run-level RTL principle above):
#   1. w:bidiVisual on the table   -> columns run right-to-left (first
#      column sits on the right).
#   2. w:bidi on each cell paragraph -> the cell's text flows/aligns right.
#   3. text is written via add_run  -> every run gets w:rtl and the David
#      (cs) font from the style.
#   4. w:tblLayout=fixed + Cm widths -> stable column widths (otherwise the
#      description column gets squeezed).
#
# OOXML is order-sensitive: children of tblPr/pPr must appear in schema
# order or Word silently drops them. _insert_ordered places an element at
# its correct position regardless of what is already there.

_TBLPR_ORDER = [qn('w:' + n) for n in (
    'tblStyle', 'tblpPr', 'tblOverlap', 'bidiVisual', 'tblStyleRowBandSize',
    'tblStyleColBandSize', 'tblW', 'jc', 'tblCellSpacing', 'tblInd',
    'tblBorders', 'shd', 'tblLayout', 'tblCellMar', 'tblLook', 'tblCaption',
    'tblDescription',
)]

_PPR_ORDER = [qn('w:' + n) for n in (
    'pStyle', 'keepNext', 'keepLines', 'pageBreakBefore', 'framePr',
    'widowControl', 'numPr', 'suppressLineNumbers', 'pBdr', 'shd', 'tabs',
    'suppressAutoHyphens', 'kinsoku', 'wordWrap', 'overflowPunct',
    'topLinePunct', 'autoSpaceDE', 'autoSpaceDN', 'bidi', 'adjustRightInd',
    'snapToGrid', 'spacing', 'ind', 'contextualSpacing', 'mirrorIndents',
    'suppressOverlap', 'jc', 'textDirection', 'textAlignment',
    'textboxTightWrap', 'outlineLvl', 'divId', 'cnfStyle', 'rPr', 'sectPr',
    'pPrChange',
)]


def _insert_ordered(parent, child, order):
    """Insert child into parent at its schema-correct position (per order)."""
    pos = order.index(child.tag)
    for existing in parent:
        try:
            if order.index(existing.tag) > pos:
                existing.addprevious(child)
                return child
        except ValueError:
            continue  # unknown/foreign element - skip it
    parent.append(child)
    return child


def _set_cell_rtl(cell, text, bold=False):
    """Write text into a single cell with correct RTL direction (paragraph + run)."""
    p = cell.paragraphs[0]
    p.clear()  # drop any existing runs; NOT p.text="" which leaves an empty run
    _insert_ordered(p._p.get_or_add_pPr(), OxmlElement('w:bidi'), _PPR_ORDER)
    # אין לקבוע כאן w:jc. ראה "עקרון היישור בתאי טבלה" בראש הקובץ.
    if bold:
        add_run(p, text.replace(BOLD_OPEN, '').replace(BOLD_CLOSE, ''), bold=True)
    else:
        add_formatted_text(p, text)


def add_table(doc, headers, rows, widths=None, header_bold=True, style='Table Grid'):
    """Add an RTL table to the document.

    headers      : list of column-header strings.
    rows         : list of rows, each a list of cell strings.
    widths       : optional column widths in cm. The first entry is the
                   right-most column (columns run right-to-left).
    header_bold  : whether the header row is bold.
    style        : table style (default: grid lines).

    Returns the table object for further tweaks.
    """
    n = len(headers)
    t = doc.add_table(rows=1, cols=n)
    t.style = style
    t.alignment = WD_TABLE_ALIGNMENT.RIGHT
    t.autofit = False                                                 # emits tblLayout=fixed
    tblPr = t._tbl.tblPr
    _insert_ordered(tblPr, OxmlElement('w:bidiVisual'), _TBLPR_ORDER)  # RTL column order

    for i, h in enumerate(headers):
        _set_cell_rtl(t.rows[0].cells[i], h, bold=header_bold)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            _set_cell_rtl(cells[i], str(val))

    if widths:
        for r in t.rows:
            for i, w in enumerate(widths):
                r.cells[i].width = Cm(w)
    return t


def save(doc, output_path):
    doc.save(str(output_path))
