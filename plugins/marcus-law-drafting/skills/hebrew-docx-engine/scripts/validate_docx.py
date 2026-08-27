#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
validate_docx.py - שער אימות לכל DOCX שנבנה דרך docx_hebrew_engine.

הסקריפט אינו בודק תוכן. הוא בודק שהמסמך מקיים את חוזה המנוע, כלומר
שהוא ייפתח ב-Word כמסמך עברי תקין עם העיצוב של התבנית.

הכללים נגזרים משלושת העקרונות הקריטיים ב-hebrew-docx-engine/SKILL.md,
ונמדדו מול פלט אמיתי של המנוע ולא הונחו:

  1. כל פסקה בסגנון בשם מתוך הסט שהמנוע פולט.
  2. כל run נושא <w:rtl/>. בלעדיו Word מתייחס לעברית כ-LTR ומציג אותה
     בפונט ה-ascii במקום David.
  3. אף run אינו קובע rFonts או sz. סגנון Normal מפריד בין השפות
     (עברית דרך szCs, אנגלית דרך sz), וקביעה ברמת ה-run שוברת את גודל
     האנגלית.
  4. אף פסקה בתוך תא טבלה אינה קובעת w:jc. בפסקת bidi הערך right ממופה
     ל-end שהוא צד שמאל, ולכן jc=right מיישר שמאלה.
  5. ה-sectPr נושא <w:bidi/>. בלעדיו הסקשן כולו LTR.
  6. כל טבלה נושאת <w:bidiVisual/>, וכל פסקת תא נושאת <w:bidi/>.
  7. אין מקף ארוך או בינוני בשום מקום בטקסט.

שים לב: פסקאות הגוף אינן נושאות w:bidi ברמת הפסקה, וזה תקין. ה-RTL שלהן
מגיע מהסגנון ומה-sectPr. רק פסקאות בתוך תאי טבלה מקבלות w:bidi מפורש.
זה נמדד על פלט המנוע; אל תוסיף כלל שדורש w:bidi בכל פסקה.

קודי יציאה: 0 תקין, 1 נמצאו הפרות, 2 שגיאת קלט.
"""

import os
import sys
from pathlib import Path

_HERE = Path(__file__).resolve().parent
if str(_HERE) not in sys.path:
    sys.path.insert(0, str(_HERE))

try:
    import docx
    from docx.oxml.ns import qn
except ImportError:  # pragma: no cover
    sys.stderr.write("חסר python-docx. התקן: pip install python-docx\n")
    raise SystemExit(2)

# מקור-האמת לשמות הסגנונות הוא המנוע עצמו, לא רשימה מקבילה כאן.
import docx_hebrew_engine as hd

ALLOWED_STYLES = {
    hd.STYLE_TITLE,
    hd.STYLE_HEADING,
    hd.STYLE_HEADING3,
    hd.STYLE_BODY,
    hd.STYLE_QUOTE,
    hd.STYLE_EXHIBIT,
    "Normal",  # add_plain
}

BAD_DASHES = {
    "—": "EM DASH",
    "–": "EN DASH",
    "―": "HORIZONTAL BAR",
}


class Finding:
    __slots__ = ("rule", "where", "detail", "excerpt")

    def __init__(self, rule, where, detail, excerpt=""):
        self.rule = rule
        self.where = where
        self.detail = detail
        self.excerpt = excerpt

    def as_dict(self):
        return {"rule": self.rule, "where": self.where,
                "detail": self.detail, "excerpt": self.excerpt}


def _excerpt(text, n=40):
    text = (text or "").strip().replace("\n", " ")
    return text[:n] + ("..." if len(text) > n else "")


def _pPr(p):
    return p._p.find(qn("w:pPr"))


def _has(el, tag):
    return el is not None and el.find(qn(tag)) is not None


def _check_runs(p, where, findings):
    for j, r in enumerate(p.runs):
        if not r.text:
            continue  # run ריק אינו נושא טקסט ואינו משפיע על התצוגה
        rPr = r._r.find(qn("w:rPr"))
        if not _has(rPr, "w:rtl"):
            findings.append(Finding(
                "run-rtl", where + " run " + str(j),
                "run בלי <w:rtl/>. Word יציג את העברית כ-LTR ובפונט האנגלי.",
                _excerpt(r.text)))
        if _has(rPr, "w:rFonts"):
            findings.append(Finding(
                "run-rfonts", where + " run " + str(j),
                "run קובע rFonts. הפונט מגיע מהסגנון; קביעה כאן שוברת את גודל האנגלית.",
                _excerpt(r.text)))
        if _has(rPr, "w:sz"):
            findings.append(Finding(
                "run-sz", where + " run " + str(j),
                "run קובע sz. הגודל מגיע מהסגנון; קביעה כאן שוברת את גודל האנגלית.",
                _excerpt(r.text)))


def _check_dashes(text, where, findings):
    for ch, name in BAD_DASHES.items():
        if ch in text:
            col = text.index(ch)
            findings.append(Finding(
                "em-dash", where,
                "נמצא {0} (U+{1:04X}) בעמודה {2}. השתמש במקף רגיל.".format(
                    name, ord(ch), col),
                _excerpt(text)))


def validate(path):
    """מחזיר רשימת Finding. רשימה ריקה = המסמך מקיים את חוזה המנוע."""
    findings = []
    d = docx.Document(path)

    # --- כלל 5: sectPr ---
    for i, section in enumerate(d.sections):
        if not _has(section._sectPr, "w:bidi"):
            findings.append(Finding(
                "section-bidi", "section " + str(i),
                "ה-sectPr אינו נושא <w:bidi/>. הסקשן כולו ייפתח כ-LTR."))

    # --- כללים 1, 2, 3, 7 על פסקאות הגוף ---
    for i, p in enumerate(d.paragraphs):
        where = "פסקה " + str(i)
        style = p.style.name if p.style is not None else None
        if style not in ALLOWED_STYLES:
            findings.append(Finding(
                "style", where,
                "סגנון {0!r} אינו מסגנונות המנוע. מותרים: {1}.".format(
                    style, ", ".join(sorted(ALLOWED_STYLES))),
                _excerpt(p.text)))
        _check_runs(p, where, findings)
        _check_dashes(p.text, where, findings)

    # --- כללים 4, 6 על טבלאות ---
    for ti, t in enumerate(d.tables):
        tblPr = t._tbl.find(qn("w:tblPr"))
        if not _has(tblPr, "w:bidiVisual"):
            findings.append(Finding(
                "table-bidivisual", "טבלה " + str(ti),
                "הטבלה אינה נושאת <w:bidiVisual/>. העמודות יזרמו משמאל לימין."))
        for ri, row in enumerate(t.rows):
            for ci, cell in enumerate(row.cells):
                for pi, p in enumerate(cell.paragraphs):
                    where = "טבלה {0} שורה {1} תא {2} פסקה {3}".format(ti, ri, ci, pi)
                    pPr = _pPr(p)
                    if not _has(pPr, "w:bidi"):
                        findings.append(Finding(
                            "cell-bidi", where,
                            "פסקת תא בלי <w:bidi/>. הטקסט בתא יזרום שמאלה.",
                            _excerpt(p.text)))
                    if _has(pPr, "w:jc"):
                        findings.append(Finding(
                            "cell-jc", where,
                            "פסקת תא קובעת w:jc. בפסקת bidi הערך right ממופה ל-end "
                            "שהוא צד שמאל, ולכן זה מיישר שמאלה. הסר את w:jc.",
                            _excerpt(p.text)))
                    _check_runs(p, where, findings)
                    _check_dashes(p.text, where, findings)

    return findings


def format_report(path, findings):
    bar = "=" * 60
    lines = [bar, "אימות מבנה DOCX: " + os.path.basename(path), bar]
    if not findings:
        lines += ["", "✓ עבר. המסמך מקיים את חוזה docx_hebrew_engine.", "", bar]
        return "\n".join(lines)

    by_rule = {}
    for f in findings:
        by_rule.setdefault(f.rule, []).append(f)

    lines += ["", "✗ נכשל. {0} הפרות ב-{1} כללים.".format(
        len(findings), len(by_rule)), ""]
    for rule in sorted(by_rule):
        group = by_rule[rule]
        lines.append("[{0}] {1} מופעים".format(rule, len(group)))
        for f in group[:10]:
            lines.append("  ✗ {0}: {1}".format(f.where, f.detail))
            if f.excerpt:
                lines.append("      -> " + repr(f.excerpt))
        if len(group) > 10:
            lines.append("  ... ועוד {0}".format(len(group) - 10))
        lines.append("")
    lines.append(bar)
    return "\n".join(lines)


def main(argv=None):
    import argparse
    import json

    for stream in (sys.stdout, sys.stderr):
        try:
            stream.reconfigure(encoding="utf-8")
        except (AttributeError, ValueError):
            pass

    ap = argparse.ArgumentParser(
        description="אימות שמסמך DOCX מקיים את חוזה docx_hebrew_engine. "
                    "קוד יציאה 0 תקין, 1 הפרות, 2 שגיאת קלט.")
    ap.add_argument("docx", nargs="+", help="נתיב לקובץ DOCX אחד או יותר")
    ap.add_argument("--json", action="store_true", help="פלט JSON")
    args = ap.parse_args(argv)

    results = {}
    worst = 0
    for path in args.docx:
        if not os.path.isfile(path):
            sys.stderr.write("לא נמצא קובץ: " + path + "\n")
            worst = max(worst, 2)
            continue
        try:
            findings = validate(path)
        except Exception as e:  # noqa: BLE001
            sys.stderr.write("לא ניתן לקרוא {0}: {1}\n".format(path, e))
            worst = max(worst, 2)
            continue
        results[path] = findings
        if findings:
            worst = max(worst, 1)

    if args.json:
        print(json.dumps(
            dict((p, [f.as_dict() for f in fs]) for p, fs in results.items()),
            ensure_ascii=False, indent=2))
    else:
        for path, findings in results.items():
            print(format_report(path, findings))

    return worst


if __name__ == "__main__":
    raise SystemExit(main())
