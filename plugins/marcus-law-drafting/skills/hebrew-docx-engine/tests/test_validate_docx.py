#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
בדיקות ל-validate_docx.py.

הבדיקה החשובה כאן היא לא שהוולידטור מאשר מסמך תקין, אלא ש**הוא נכשל
כשצריך**. ולידטור שתמיד מחזיר 0 גרוע מאין ולידטור: הוא מייצר ראיה שקרית
לכך שהמסמך נבדק.

לכן לכל אחד מתשעת הכללים יש כאן מקרה שמפר אותו במפורש, והבדיקה דורשת
גם קוד יציאה 1 וגם שהכלל הנכון הוא זה שדווח.

הרצה:  python3 -m unittest discover -s tests
   או:  python3 tests/test_validate_docx.py
"""

import sys
import tempfile
import unittest
from pathlib import Path

_SCRIPTS = Path(__file__).resolve().parent.parent / "scripts"
sys.path.insert(0, str(_SCRIPTS))

import docx_hebrew_engine as hd  # noqa: E402
import validate_docx as vd  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402


def _base_doc():
    """מסמך תקין מינימלי שנבנה דרך המנוע: פסקה, כותרת וטבלה."""
    doc = hd.open_document()
    hd.add_title(doc, "החלטה")
    hd.add_clause(doc, "פסקה עברית עם English inside.")
    hd.add_table(doc, headers=["#", "שם"], rows=[["1", "כתב ערר"]],
                 widths=[0.8, 4.6])
    return doc


def _rules(doc):
    """שומר את המסמך לקובץ זמני ומחזיר את סט הכללים שהופרו."""
    with tempfile.TemporaryDirectory() as tmp:
        path = str(Path(tmp) / "t.docx")
        hd.save(doc, path)
        return {f.rule for f in vd.validate(path)}


class TestValidDocument(unittest.TestCase):
    def test_engine_output_passes(self):
        """פלט נקי של המנוע חייב לעבור בלי אף ממצא."""
        self.assertEqual(_rules(_base_doc()), set())

    def test_body_paragraphs_need_no_bidi(self):
        """
        פסקאות גוף אינן נושאות w:bidi ברמת הפסקה, וזה תקין.
        הבדיקה הזו נועלת את ההתנהגות כדי שאיש לא יוסיף כלל שדורש w:bidi
        בכל פסקה, מה שהיה מכשיל כל מסמך תקין במערכת.
        """
        doc = _base_doc()
        pPr = doc.paragraphs[1]._p.find(qn("w:pPr"))
        has_bidi = pPr is not None and pPr.find(qn("w:bidi")) is not None
        self.assertFalse(has_bidi, "המנוע שינה התנהגות: עדכן את הוולידטור")
        self.assertEqual(_rules(doc), set())


class TestEachRuleFires(unittest.TestCase):
    """לכל כלל, מסמך שמפר אותו בכוונה חייב להיתפס."""

    def _assert_catches(self, rule, mutate):
        rules = _rules(mutate(_base_doc()))
        self.assertIn(rule, rules,
                      "הכלל {0!r} לא נתפס. דווחו: {1}".format(rule, rules or "כלום"))

    def test_run_without_rtl(self):
        def mutate(doc):
            rPr = doc.paragraphs[1].runs[0]._r.find(qn("w:rPr"))
            rPr.remove(rPr.find(qn("w:rtl")))
            return doc
        self._assert_catches("run-rtl", mutate)

    def test_run_sets_rfonts(self):
        def mutate(doc):
            rPr = doc.paragraphs[1].runs[0]._r.get_or_add_rPr()
            el = OxmlElement("w:rFonts")
            el.set(qn("w:ascii"), "Arial")
            rPr.append(el)
            return doc
        self._assert_catches("run-rfonts", mutate)

    def test_run_sets_sz(self):
        def mutate(doc):
            rPr = doc.paragraphs[1].runs[0]._r.get_or_add_rPr()
            el = OxmlElement("w:sz")
            el.set(qn("w:val"), "28")
            rPr.append(el)
            return doc
        self._assert_catches("run-sz", mutate)

    def test_cell_paragraph_sets_jc(self):
        """המלכודת המתועדת: jc=right בפסקת bidi מיישר שמאלה."""
        def mutate(doc):
            p = doc.tables[0].rows[0].cells[0].paragraphs[0]
            el = OxmlElement("w:jc")
            el.set(qn("w:val"), "right")
            p._p.get_or_add_pPr().append(el)
            return doc
        self._assert_catches("cell-jc", mutate)

    def test_cell_paragraph_without_bidi(self):
        def mutate(doc):
            pPr = doc.tables[0].rows[0].cells[0].paragraphs[0]._p.find(qn("w:pPr"))
            pPr.remove(pPr.find(qn("w:bidi")))
            return doc
        self._assert_catches("cell-bidi", mutate)

    def test_table_without_bidivisual(self):
        def mutate(doc):
            tblPr = doc.tables[0]._tbl.find(qn("w:tblPr"))
            tblPr.remove(tblPr.find(qn("w:bidiVisual")))
            return doc
        self._assert_catches("table-bidivisual", mutate)

    def test_section_without_bidi(self):
        def mutate(doc):
            sect = doc.sections[0]._sectPr
            sect.remove(sect.find(qn("w:bidi")))
            return doc
        self._assert_catches("section-bidi", mutate)

    def test_foreign_style(self):
        def mutate(doc):
            doc.add_paragraph("פסקה בסגנון זר", style="List Paragraph")
            return doc
        self._assert_catches("style", mutate)

    def test_em_dash(self):
        """
        המנוע מנקה מקפים ארוכים ב-add_run, ולכן ההפרה מוזרקת אחרי הבנייה.
        זה מדמה מסמך שנערך בכלי אחר אחרי שיצא מהמנוע.
        """
        def mutate(doc):
            doc.paragraphs[1].runs[0].text = "טקסט עם מקף ארוך — כאן"
            return doc
        self._assert_catches("em-dash", mutate)


class TestSanitizerStillGuards(unittest.TestCase):
    def test_engine_strips_em_dash_on_the_way_in(self):
        """
        רשת הביטחון של המנוע: טקסט עם מקף ארוך שנכנס דרך add_clause יוצא נקי,
        ולכן הוולידטור לא ידווח em-dash. אם זה נשבר, שתי ההגנות נשברו יחד.
        """
        doc = hd.open_document()
        hd.add_clause(doc, "טקסט עם מקף ארוך — כאן")
        self.assertNotIn("em-dash", _rules(doc))


if __name__ == "__main__":
    unittest.main(verbosity=2)
