#!/usr/bin/env python3
"""
Pre-process a legal pleading DOCX before LibreOffice conversion.

Two operations:

1. **Line spacing**
   - Body paragraphs (outside tables) → 1.5 lines (court-readable spacing).
   - Paragraphs inside tables → 1.15 lines + 0pt before/after (tables stay
     compact: header table, signature block, etc., shouldn't blow up to
     1.5-line spacing).
   This matches what Israeli civil-procedure standard expects: airy body
   text, tight tables.

2. **Em-dash sanitisation**
   - Every paragraph and run text replaces em dash (U+2014) and en dash
     (U+2013) with a plain hyphen-minus. AI-generated drafts often contain
     em dashes that betray non-human authorship; courts and practitioners
     prefer the look of a Word-typed document.

Public API:
    prepare_pleading(docx_in, docx_out,
                     body_spacing=1.5, table_spacing=1.15,
                     sanitize_dashes=True)
"""

import os
import sys
from typing import Optional

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt

# Local import — works when this script is on sys.path
try:
    from sanitize_text import sanitize as _sanitize_text
except ImportError:
    # Fallback — manual replace
    def _sanitize_text(s):
        if not isinstance(s, str):
            return s
        return s.replace('—', '-').replace('–', '-').replace('―', '-')


def _apply_spacing(p, multiple: float, *, zero_padding: bool = False):
    pf = p.paragraph_format
    pf.line_spacing = multiple
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    if zero_padding:
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)


def _sanitize_paragraph_runs(p):
    """Replace banned dashes inside every run of the paragraph (preserves run-level formatting)."""
    for run in p.runs:
        if run.text:
            new = _sanitize_text(run.text)
            if new != run.text:
                run.text = new


def _walk_tables(tables, table_spacing: float, sanitize_dashes: bool, counters: dict):
    for tbl in tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _apply_spacing(p, table_spacing, zero_padding=True)
                    if sanitize_dashes:
                        _sanitize_paragraph_runs(p)
                    counters['table'] += 1
                # Recurse into nested tables
                if cell.tables:
                    _walk_tables(cell.tables, table_spacing, sanitize_dashes, counters)


def prepare_pleading(docx_in: str,
                     docx_out: str,
                     body_spacing: float = 1.5,
                     table_spacing: float = 1.15,
                     sanitize_dashes: bool = True) -> str:
    """
    Apply spacing rules and sanitise text in a pleading DOCX.

    Args:
        docx_in: Path to source DOCX.
        docx_out: Where to save the prepared DOCX.
        body_spacing: Line spacing multiple for body-level paragraphs (default 1.5).
        table_spacing: Line spacing multiple for paragraphs inside tables (default 1.15).
        sanitize_dashes: Replace em/en dashes with hyphens (default True).

    Returns:
        docx_out (the path written to).
    """
    doc = Document(docx_in)
    counters = {'body': 0, 'table': 0}

    # Body-level paragraphs (NOT including paragraphs inside table cells)
    for p in doc.paragraphs:
        _apply_spacing(p, body_spacing)
        if sanitize_dashes:
            _sanitize_paragraph_runs(p)
        counters['body'] += 1

    # Table cells (recurses into nested tables)
    _walk_tables(doc.tables, table_spacing, sanitize_dashes, counters)

    doc.save(docx_out)
    return docx_out


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage: prepare_docx.py INPUT.docx OUTPUT.docx [body_spacing] [table_spacing]")
        sys.exit(1)
    src, dst = sys.argv[1], sys.argv[2]
    body = float(sys.argv[3]) if len(sys.argv) > 3 else 1.5
    table = float(sys.argv[4]) if len(sys.argv) > 4 else 1.15
    out = prepare_pleading(src, dst, body_spacing=body, table_spacing=table)
    print(f"Prepared: {out}")
    print(f"  body spacing:  {body}")
    print(f"  table spacing: {table}")
