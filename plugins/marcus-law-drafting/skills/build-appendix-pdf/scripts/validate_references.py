#!/usr/bin/env python3
"""Extract and validate appendix references in pleading text."""

import os
import re
from typing import List, Dict, Any, Tuple

# Comprehensive reference pattern families
REFERENCE_PATTERNS = [
    # Family A: Full "נספח X" patterns with various prefixes
    r'מצ"ב\s*ומסומ[נן](?:ת|ים|ות)?\s*(?:יחדיו\s*)?נספח\s*([א-ת]{1,2}[\'])',
    r'מצ"ב\s*ומסומ[נן](?:ת|ים|ות)?\s*(?:יחדיו\s*)?נספח\s*(\d+)',
    r'מסומ[נן](?:ת|ים|ות)?\s*כנספח\s*([א-ת]{1,2}[\'])',
    r'מסומ[נן](?:ת|ים|ות)?\s*כנספח\s*(\d+)',
    r'ומסומ[נן](?:ת|ים|ות)?\s*כנספח\s*([א-ת]{1,2}[\'])',
    r'ומסומ[נן](?:ת|ים|ות)?\s*כנספח\s*(\d+)',
    r'ראו\s*נספח\s*([א-ת]{1,2}[\'])',
    r'ראו\s*נספח\s*(\d+)',
    r'ר[\']\s*נספח\s*([א-ת]{1,2}[\'])',
    r'נספח\s+([א-ת]{1,2}[\'])\s+הנ"ל',
    r'נספח\s+(\d+)\s+הנ"ל',
    r'בנספח\s*([א-ת]{1,2}[\'])',
    r'בנספח\s*(\d+)',
    
    # Family B: "ומסומן X'" without "נספח" prefix (Gilad Cohen style)
    r'מצ"ב\s*ומסומ[נן](?:ת|ים|ות)?\s*(?:יחדיו\s*)?([א-ת]{1,2}[\'])',
    r'ומסומ[נן](?:ת|ים|ות)?\s*(?:יחדיו\s*)?([א-ת]{1,2}[\'])',
    
    # Family C: Catch-all for mangled PDF text
    r'נספח\s*([א-ת]{1,2}[\'])',
    r'נספח\s*(\d+)',
]

FALSE_POSITIVES = {"עמ'", "עמ", "תוכן", "תוכן'"}


def extract_references(text: str) -> List[Dict[str, Any]]:
    """
    Extract all appendix references from pleading text.
    
    Uses three families of regex patterns (A: full "נספח X", B: Gilad Cohen style,
    C: catch-all) and deduplicates by (reference_id, line_number).
    
    Args:
        text: Full pleading text to analyze
    
    Returns:
        List of dicts with keys:
        - id: Appendix identifier (א' or 1)
        - line_num: Line number where found (1-based)
        - context: Text surrounding the reference
        - pattern_family: 'A', 'B', or 'C'
    """
    references = []
    seen = set()
    
    lines = text.split('\n')
    
    for line_idx, line in enumerate(lines, 1):
        for pattern_idx, pattern in enumerate(REFERENCE_PATTERNS):
            matches = re.finditer(pattern, line)
            for match in matches:
                ref_id = match.group(1).strip()
                
                # Skip false positives
                if ref_id in FALSE_POSITIVES:
                    continue
                
                # Determine pattern family
                if pattern_idx < 13:
                    family = 'A'
                elif pattern_idx < 15:
                    family = 'B'
                else:
                    family = 'C'
                
                # Deduplicate
                key = (ref_id, line_idx)
                if key in seen:
                    continue
                seen.add(key)
                
                # Extract context (50 chars before and after, or to line boundaries)
                start = max(0, match.start() - 50)
                end = min(len(line), match.end() + 50)
                context = line[start:end].strip()
                
                references.append({
                    'id': ref_id,
                    'line_num': line_idx,
                    'context': context,
                    'pattern_family': family
                })
    
    return references


def validate(text: str, appendix_list: List[Dict[str, Any]]) -> Dict[str, Any]:
    """
    Validate appendix references against appendix list.
    
    Args:
        text: Full pleading text
        appendix_list: List of dicts with 'id' and 'name' keys
    
    Returns:
        Dict with:
        - is_valid: Boolean (True if all references valid and all appendices referenced)
        - unique_appendices: Set of unique appendix IDs found
        - missing_references: List of appendix IDs in list but not referenced
        - orphaned_references: List of appendix IDs referenced but not in list
        - reference_count: Total references found
        - errors: List of error strings
        - warnings: List of warning strings
    """
    references = extract_references(text)
    referenced_ids = {ref['id'] for ref in references}
    listed_ids = {app['id'] for app in appendix_list}
    
    missing_refs = listed_ids - referenced_ids
    orphaned_refs = referenced_ids - listed_ids
    
    errors = []
    warnings = []
    
    if orphaned_refs:
        errors.append(f"Found references to appendices not in list: {sorted(orphaned_refs)}")
    
    if missing_refs:
        errors.append(f"Appendices in list but not referenced in pleading: {sorted(missing_refs)}")

    is_valid = len(errors) == 0
    
    return {
        'is_valid': is_valid,
        'unique_appendices': sorted(referenced_ids),
        'missing_references': sorted(missing_refs),
        'orphaned_references': sorted(orphaned_refs),
        'reference_count': len(references),
        'errors': errors,
        'warnings': warnings,
        'references': references
    }


def format_validation_report(validation_result: Dict[str, Any]) -> str:
    """
    Format validation result as human-readable report.
    
    Args:
        validation_result: Result from validate() function
    
    Returns:
        Formatted report string
    """
    lines = []
    lines.append("=" * 60)
    lines.append("APPENDIX VALIDATION REPORT")
    lines.append("=" * 60)
    
    lines.append(f"\nStatus: {'✓ VALID' if validation_result['is_valid'] else '✗ INVALID'}")
    lines.append(f"References found: {validation_result['reference_count']}")
    lines.append(f"Unique appendices: {', '.join(validation_result['unique_appendices']) or 'none'}")
    
    if validation_result['errors']:
        lines.append("\nERRORS:")
        for error in validation_result['errors']:
            lines.append(f"  ✗ {error}")
    
    if validation_result['warnings']:
        lines.append("\nWARNINGS:")
        for warning in validation_result['warnings']:
            lines.append(f"  ⚠ {warning}")
    
    if validation_result['references']:
        lines.append("\nREFERENCES FOUND:")
        for ref in sorted(validation_result['references'], key=lambda r: r['line_num']):
            lines.append(f"  Line {ref['line_num']:3d}: {ref['id']:5s} (Family {ref['pattern_family']}) - {ref['context'][:50]}")
    
    lines.append("\n" + "=" * 60)
    
    return "\n".join(lines)


def extract_pleading_text(path: str) -> str:
    """
    Read the pleading body as plain text, from either DOCX or PDF.

    Detection is by magic bytes, not extension, matching Iron Rule 1 -- a DOCX
    saved as .pdf (or the reverse) is common in the wild and must not silently
    produce empty text, which would make every reference look missing.
    """
    with open(path, 'rb') as fh:
        head = fh.read(4)

    if head[:2] == b'PK':  # zip container -> DOCX
        from docx import Document
        doc = Document(path)
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.extend(p.text for p in cell.paragraphs)
        return '\n'.join(parts)

    if head == b'%PDF':
        import fitz
        doc = fitz.open(path)
        try:
            return '\n'.join(page.get_text() for page in doc)
        finally:
            doc.close()

    # Neither container: treat as plain text rather than failing, so a caller
    # that already extracted the text can pass a .txt file.
    with open(path, encoding='utf-8', errors='replace') as fh:
        return fh.read()


def _self_test() -> int:
    test_text = """
    כמפורט בנספח א' להלן.
    ראו נספח ב' לפרטים נוספים.
    המסמכים ומסומנים כנספח ג'.
    בנספח ד' מוצגת הרשימה המלאה.
    """
    test_list = [
        {'id': "א'", 'name': 'Document 1'},
        {'id': "ב'", 'name': 'Document 2'},
        {'id': "ג'", 'name': 'Document 3'},
        {'id': "ד'", 'name': 'Document 4'},
    ]
    result = validate(test_text, test_list)
    print(format_validation_report(result))
    return 0 if result['is_valid'] else 1


def main(argv=None) -> int:
    import argparse
    import json
    import sys

    # הדוח מכיל עברית וסימני ✓/✗. קונסולת Windows מוגדרת ל-cp1255 כברירת מחדל
    # ותקרוס עליהם, ולכן כופים UTF-8 על הפלט לפני כל הדפסה.
    for stream in (sys.stdout, sys.stderr):
        try:
            stream.reconfigure(encoding='utf-8')
        except (AttributeError, ValueError):
            pass

    ap = argparse.ArgumentParser(
        description="אימות הפניות צולבות לנספחים מול גוף כתב הטענות. "
                    "קוד יציאה 0 = כל ההפניות תקינות, 1 = נמצאו שגיאות.")
    ap.add_argument('pleading', nargs='?',
                    help='נתיב לכתב הטענות (DOCX / PDF / TXT)')
    ap.add_argument('--appendix-list', dest='appendix_list',
                    help='נתיב לקובץ JSON עם [{"id": "1", "name": "..."}] '
                         'או ה-JSON עצמו כמחרוזת')
    ap.add_argument('--json', action='store_true',
                    help='פלט JSON במקום דוח קריא')
    ap.add_argument('--self-test', action='store_true',
                    help='הרצת בדיקה עצמית על טקסט לדוגמה')
    args = ap.parse_args(argv)

    if args.self_test:
        return _self_test()

    if not args.pleading or not args.appendix_list:
        ap.error('נדרשים גם pleading וגם --appendix-list (או --self-test)')

    raw = args.appendix_list
    if os.path.isfile(raw):
        with open(raw, encoding='utf-8') as fh:
            appendix_list = json.load(fh)
    else:
        appendix_list = json.loads(raw)

    text = extract_pleading_text(args.pleading)
    if not text.strip():
        print('שגיאה: לא חולץ טקסט מכתב הטענות. '
              'ייתכן שזה PDF סרוק ללא שכבת טקסט: הרץ OCR תחילה.')
        return 2

    result = validate(text, appendix_list)
    if args.json:
        printable = {k: v for k, v in result.items() if k != 'references'}
        print(json.dumps(printable, ensure_ascii=False, indent=2))
    else:
        print(format_validation_report(result))
    return 0 if result['is_valid'] else 1


if __name__ == '__main__':
    import sys
    sys.exit(main())
