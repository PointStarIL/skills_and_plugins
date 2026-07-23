#!/usr/bin/env python3
"""Build a Hebrew (RTL) Word document of a party-claims analysis.

Input  : a JSON file describing the extracted claims (schema below).
Output : a .docx file.

Usage:
    python build_claims_docx.py <analysis.json> <output.docx>

JSON schema
-----------
{
  "case_ref": "ערר <מספר> — <נושא התיק>",                # required, appears as subtitle
  "subtitle": "optional extra subtitle line",
  "parties": [
    {
      "name": "העוררים",                                   # required
      "threshold_claims": [                                # optional (טענות סף)
        {"heading": "שיהוי", "text": "...", "source": "כתב ערר, עמ' 4"}
      ],
      "argument_heads": [                                  # required list
        {
          "heading": "ראש טיעון 1 — סמכות הוועדה",
          "claims": [
            {"text": "העוררים טוענים כי...", "source": "כתב ערר, ס' 12"}
          ]
        }
      ]
    }
  ]
}

`source` is optional on every claim. Text is emitted verbatim — the model that
produces the JSON is responsible for neutrality and thematic ordering (see the
bundled claims-extraction-guide.md).
"""
import json
import sys

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor
except ImportError:
    sys.stderr.write(
        "python-docx is required. Install it once with:\n"
        "    pip install python-docx\n"
    )
    sys.exit(2)

BODY_FONT = "David"
GREY = RGBColor(0x6B, 0x72, 0x80)


def _rtl(paragraph):
    """Mark a paragraph right-to-left and right-aligned."""
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    return paragraph


def _set_run_rtl(run):
    """Set the complex-script (rtl) flag on a run so Hebrew renders correctly."""
    rPr = run._r.get_or_add_rPr()
    rtl = OxmlElement("w:rtl")
    rtl.set(qn("w:val"), "1")
    rPr.append(rtl)


def _para(doc, text="", *, bold=False, size=None, color=None, style=None,
          space_before=None, space_after=None):
    p = doc.add_paragraph(style=style)
    _rtl(p)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.font.name = BODY_FONT
        if size is not None:
            run.font.size = Pt(size)
        if color is not None:
            run.font.color.rgb = color
        _set_run_rtl(run)
    return p


def _claim(doc, claim):
    text = (claim.get("text") or "").strip()
    if not text:
        return
    p = _para(doc, space_after=6)
    run = p.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = Pt(12)
    _set_run_rtl(run)
    src = (claim.get("source") or "").strip()
    if src:
        p.add_run("  ")
        s = p.add_run(f"({src})")
        s.italic = True
        s.font.size = Pt(10)
        s.font.color.rgb = GREY
        s.font.name = BODY_FONT
        _set_run_rtl(s)


def build(data, out_path):
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(12)

    _para(doc, "תמצית טענות הצדדים", bold=True, size=18, space_after=4)
    case_ref = (data.get("case_ref") or "").strip()
    if case_ref:
        _para(doc, case_ref, bold=True, size=13, color=GREY, space_after=2)
    subtitle = (data.get("subtitle") or "").strip()
    if subtitle:
        _para(doc, subtitle, size=11, color=GREY, space_after=8)

    parties = data.get("parties") or []
    if not parties:
        _para(doc, "לא נמצאו טענות.", color=GREY)

    for party in parties:
        name = (party.get("name") or "צד").strip()
        h = doc.add_heading(level=1)
        _rtl(h)
        r = h.add_run(name)
        _set_run_rtl(r)

        threshold = party.get("threshold_claims") or []
        if threshold:
            sh = doc.add_heading(level=2)
            _rtl(sh)
            r = sh.add_run("טענות סף")
            _set_run_rtl(r)
            for tc in threshold:
                head = (tc.get("heading") or "").strip()
                if head:
                    _para(doc, head, bold=True, size=12, space_before=4,
                          space_after=2)
                _claim(doc, tc)

        heads = party.get("argument_heads") or []
        if heads:
            sh = doc.add_heading(level=2)
            _rtl(sh)
            r = sh.add_run("טענות לגופו של עניין")
            _set_run_rtl(r)
        for head in heads:
            head_title = (head.get("heading") or "").strip()
            if head_title:
                _para(doc, head_title, bold=True, size=12.5,
                      space_before=6, space_after=3)
            for claim in head.get("claims") or []:
                _claim(doc, claim)

    doc.save(out_path)


def main():
    if len(sys.argv) != 3:
        sys.stderr.write(__doc__)
        sys.exit(1)
    in_path, out_path = sys.argv[1], sys.argv[2]
    with open(in_path, encoding="utf-8") as f:
        data = json.load(f)
    build(data, out_path)
    print(f"wrote {out_path}")


if __name__ == "__main__":
    main()
