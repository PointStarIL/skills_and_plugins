#!/usr/bin/env python3
"""
docx_rtl_helper.py

ספריית עזר ליצירת DOCX בעברית RTL מלא.
שימוש: import docx_rtl_helper as rtl ואז rtl.create_doc(), rtl.add_para(), וכו'.

כל הפונקציות מוודאות שכל אובייקט במסמך מוגדר ב-RTL:
- bidi בפסקאות
- bidiVisual בטבלאות
- bidi בסקציה
- פונט עברי תקני (David) ב-ascii, hAnsi, cs
- גודל מוגדר ב-sz וגם ב-szCs
- יישור נכון (RIGHT, JUSTIFY, CENTER)
- מרווח שורות 1.5 לגוף
- שוליים 2.54 ס"מ
"""

from copy import deepcopy

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ============== הגדרות ברירת מחדל ==============

DEFAULT_FONT = "David"
DEFAULT_SIZE = 12
DEFAULT_LINE_SPACING = 1.5
DEFAULT_MARGIN_CM = 2.54


# ============== פונקציות בסיס ==============

def set_rtl_paragraph(paragraph):
    """מוסיף w:bidi לפסקה כדי להגדירה כ-RTL."""
    pPr = paragraph._p.get_or_add_pPr()
    existing = pPr.find(qn('w:bidi'))
    if existing is None:
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        pPr.append(bidi)


def set_rtl_section(section):
    """מוסיף w:bidi לסקציה (להשפעה על מספור עמודים, וכו')."""
    sectPr = section._sectPr
    existing = sectPr.find(qn('w:bidi'))
    if existing is None:
        bidi = OxmlElement('w:bidi')
        sectPr.append(bidi)


def set_rtl_table(table):
    """מגדיר טבלה כ-RTL (bidiVisual + bidi לכל הפסקאות בתאים)."""
    tblPr = table._element.tblPr
    bidiVisual = tblPr.find(qn('w:bidiVisual'))
    if bidiVisual is None:
        bidiVisual = OxmlElement('w:bidiVisual')
        bidiVisual.set(qn('w:val'), '1')
        tblPr.append(bidiVisual)
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                set_rtl_paragraph(p)


def set_run_font(run, font_name=DEFAULT_FONT, size_pt=DEFAULT_SIZE, bold=False, italic=False, color=None):
    """מגדיר פונט עברי תקני ב-3 מקומות (ascii, hAnsi, cs) וגודל ב-2 מקומות (sz, szCs)."""
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color

    rPr = run._element.get_or_add_rPr()

    # פונט
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rFonts.set(qn('w:hint'), 'cs')

    # גודל ב-2 מקומות
    for tag in ['w:sz', 'w:szCs']:
        el = rPr.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            rPr.append(el)
        el.set(qn('w:val'), str(int(size_pt * 2)))

    # bold ב-2 מקומות
    if bold:
        for tag in ['w:b', 'w:bCs']:
            el = rPr.find(qn(tag))
            if el is None:
                el = OxmlElement(tag)
                rPr.append(el)
    else:
        for tag in ['w:b', 'w:bCs']:
            el = rPr.find(qn(tag))
            if el is not None:
                rPr.remove(el)


# ============== פונקציות יצירת מסמך ==============

def create_doc(margin_cm=DEFAULT_MARGIN_CM):
    """יוצר מסמך DOCX חדש עם הגדרות RTL מלאות."""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(margin_cm)
        section.bottom_margin = Cm(margin_cm)
        section.left_margin = Cm(margin_cm)
        section.right_margin = Cm(margin_cm)
        set_rtl_section(section)
    return doc


def add_para(doc, text, *,
             font=DEFAULT_FONT,
             size=DEFAULT_SIZE,
             bold=False,
             italic=False,
             align='justify',
             spacing=DEFAULT_LINE_SPACING,
             space_before=0,
             space_after=6,
             first_line_indent_cm=0.0,
             color=None):
    """מוסיף פסקה למסמך עם כל הגדרות ה-RTL הנדרשות."""
    p = doc.add_paragraph()
    set_rtl_paragraph(p)

    # יישור
    align_map = {
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
        'left': WD_ALIGN_PARAGRAPH.LEFT,
    }
    p.alignment = align_map.get(align, WD_ALIGN_PARAGRAPH.JUSTIFY)

    # פורמט פסקה
    pf = p.paragraph_format
    pf.line_spacing = spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if first_line_indent_cm:
        pf.first_line_indent = Cm(first_line_indent_cm)

    # ריצה
    if text:
        r = p.add_run(text)
        set_run_font(r, font, size, bold, italic, color)

    return p


# ============== פונקציות שכבת מסמך (כותרות, ציטוטים) ==============

def add_main_title(doc, text):
    """כותרת ראשית של המסמך, ממורכזת, פונט David מודגש 16."""
    return add_para(doc, text,
                    size=16, bold=True, align='center', spacing=1.5, space_after=12)


def add_section_heading(doc, text):
    """כותרת בלוק / סעיף, מימין, פונט David מודגש 13."""
    return add_para(doc, text,
                    size=13, bold=True, align='right', spacing=1.5,
                    space_before=12, space_after=6)


def add_subheading(doc, text):
    """כותרת משנה, מימין, פונט David מודגש 12."""
    return add_para(doc, text,
                    size=12, bold=True, align='right', spacing=1.5,
                    space_before=6, space_after=4)


def add_body(doc, text, *, bold=False):
    """פסקת גוף סטנדרטית, פונט David 12, justify, מרווח 1.5."""
    return add_para(doc, text, bold=bold)


def add_quote(doc, text):
    """ציטוט פסיקה, פונט David 11, justify, מרווח 1.15, עם הזחה."""
    p = add_para(doc, text,
                 size=11, italic=False, align='justify',
                 spacing=1.15, space_after=6,
                 first_line_indent_cm=0.0)
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.right_indent = Cm(1.0)
    return p


def add_footnote_style(doc, text):
    """הערה / הערת שוליים, פונט David 10, מרווח 1.0, ביישור ימני."""
    return add_para(doc, text,
                    size=10, align='right', spacing=1.0,
                    space_after=4, color=RGBColor(0x80, 0x80, 0x80))


# ============== פונקציות עריכת DOCX קיים ==============

def copy_paragraph_style(source_paragraph, target_paragraph):
    """מעתיק את כל מאפייני pPr (formatting) מפסקה אחת לאחרת."""
    source_pPr = source_paragraph._p.find(qn('w:pPr'))
    if source_pPr is None:
        return
    target_pPr = deepcopy(source_pPr)
    existing = target_paragraph._p.find(qn('w:pPr'))
    if existing is not None:
        target_paragraph._p.remove(existing)
    target_paragraph._p.insert(0, target_pPr)


def replace_paragraph_text(paragraph, new_text, font=DEFAULT_FONT, size=DEFAULT_SIZE, bold=False):
    """מחליף את הטקסט בפסקה קיימת, שומר על כל הפורמט."""
    # מנקה את כל הריצות הקיימות
    for run in paragraph.runs:
        run._element.getparent().remove(run._element)
    # מוסיף ריצה חדשה
    new_run = paragraph.add_run(new_text)
    set_run_font(new_run, font, size, bold)


def find_paragraphs_with_pattern(doc, pattern_regex):
    """מחזיר רשימת פסקאות שמכילות את ה-pattern."""
    import re
    pattern = re.compile(pattern_regex)
    return [(i, p) for i, p in enumerate(doc.paragraphs) if pattern.search(p.text)]


# ============== פונקציית בדיקת איכות ==============

def audit_rtl(doc):
    """סורק את המסמך ומחזיר רשימת בעיות RTL."""
    issues = []
    for i, section in enumerate(doc.sections):
        sectPr = section._sectPr
        if sectPr.find(qn('w:bidi')) is None:
            issues.append(f'סקציה {i}: חסר w:bidi')

    for i, p in enumerate(doc.paragraphs):
        if not p.text.strip():
            continue
        pPr = p._p.find(qn('w:pPr'))
        if pPr is None or pPr.find(qn('w:bidi')) is None:
            issues.append(f'פסקה {i}: חסר w:bidi (טקסט: {p.text[:30]}...)')

        # בדיקת פונט
        for run in p.runs:
            rPr = run._element.find(qn('w:rPr'))
            if rPr is None:
                issues.append(f'פסקה {i}: ריצה ללא rPr')
                continue
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                issues.append(f'פסקה {i}: ריצה ללא rFonts')
                continue
            cs = rFonts.get(qn('w:cs'))
            if not cs or cs not in ('David', 'FrankRuehl', 'Miriam', 'Arial', 'Narkisim'):
                issues.append(f'פסקה {i}: פונט CS לא תקני: {cs}')
            sz = rPr.find(qn('w:sz'))
            szCs = rPr.find(qn('w:szCs'))
            if sz is None or szCs is None:
                issues.append(f'פסקה {i}: חסר sz או szCs')

    for ti, table in enumerate(doc.tables):
        tblPr = table._element.tblPr
        if tblPr.find(qn('w:bidiVisual')) is None:
            issues.append(f'טבלה {ti}: חסר w:bidiVisual')

    return issues


# ============== נקודת כניסה לדוגמה ==============

if __name__ == '__main__':
    import sys

    if len(sys.argv) > 1 and sys.argv[1] == '--audit':
        if len(sys.argv) < 3:
            print('usage: docx_rtl_helper.py --audit /path/to/file.docx')
            sys.exit(1)
        doc = Document(sys.argv[2])
        issues = audit_rtl(doc)
        if issues:
            print(f'נמצאו {len(issues)} בעיות RTL:')
            for issue in issues:
                print(f'  - {issue}')
            sys.exit(2)
        print('עבר את בדיקת ה-RTL בהצלחה: כל הפסקאות, הסקציות, והטבלאות מוגדרות נכון.')
        sys.exit(0)

    # דוגמה ליצירת מסמך לדוגמה
    doc = create_doc()
    add_main_title(doc, 'דוגמה למסמך RTL מלא')
    add_section_heading(doc, 'כותרת בלוק לדוגמה')
    add_body(doc, 'זוהי פסקת גוף סטנדרטית. הטקסט יזרום מימין לשמאל, במרווח 1.5, פונט David גודל 12, ויישור דו-צדדי.')
    add_subheading(doc, 'כותרת משנה')
    add_body(doc, 'פסקה נוספת תחת כותרת המשנה.')
    add_quote(doc, '"זהו ציטוט פסיקה לדוגמה. הציטוט מוצג בפונט David 11, יישור דו-צדדי, ועם הזחות מימין ומשמאל."')
    add_body(doc, 'פסקה מסכמת שמראה כיצד נראה התוכן אחרי הציטוט.')

    output = '/tmp/rtl_demo.docx'
    doc.save(output)
    print(f'נשמר: {output}')
    issues = audit_rtl(Document(output))
    if issues:
        print(f'!!! נמצאו {len(issues)} בעיות:')
        for issue in issues:
            print(f'  - {issue}')
    else:
        print('בדיקת RTL: עבר.')
