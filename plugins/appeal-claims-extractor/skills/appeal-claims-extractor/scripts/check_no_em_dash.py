#!/usr/bin/env python3
"""
check_no_em_dash.py

בדיקה אוטומטית שאין מקפים ארוכים (em-dash) בטקסט פלט.
מקף ארוך הוא חתימה אופיינית של כלי AI ופוגע באמינות מסמך משפטי.

שימוש:
    python check_no_em_dash.py path/to/file.md
    python check_no_em_dash.py path/to/file.docx
    echo "טקסט לבדיקה" | python check_no_em_dash.py

קודי יציאה:
    0 = הטקסט נקי (אין מקפים ארוכים)
    1 = נמצאו מקפים ארוכים (פירוט ב-stderr)
    2 = שגיאת קלט
"""

import sys
from pathlib import Path

# התווים הבעייתיים שאסור שיופיעו בטקסט
FORBIDDEN_CHARS = {
    "—": "EM DASH (—)",
    "–": "EN DASH (–)",
    "―": "HORIZONTAL BAR (―)",
}

# חלופות מומלצות
SUGGESTIONS = """
חלופות מומלצות במקום מקף ארוך:
    1. נקודתיים ":" כשמתחיל פירוט.
    2. פסיק "," לרצף רגיל.
    3. נקודה "." ופתיחת משפט חדש.
    4. סוגריים "( )" לתוספת הסבר משנית.
    5. מילים: "כלומר", "דהיינו", "שכן", "מאחר", "כי", "וזאת".
    6. מקף קצר רגיל "-" לחיבור מילים או טווחים מספריים.
"""


def read_text_from_path(path: Path) -> str:
    """קורא טקסט מקובץ. תומך ב-txt, md, py, וכן docx (דרך python-docx אם מותקן)."""
    suffix = path.suffix.lower()

    if suffix in (".txt", ".md", ".py", ".json", ".html"):
        return path.read_text(encoding="utf-8")

    if suffix == ".docx":
        try:
            from docx import Document
        except ImportError:
            print(
                "ERROR: python-docx not installed. Install with: pip install python-docx",
                file=sys.stderr,
            )
            sys.exit(2)

        doc = Document(str(path))
        parts = []
        for para in doc.paragraphs:
            parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(parts)

    # ברירת מחדל: נסה לקרוא כטקסט
    return path.read_text(encoding="utf-8")


def find_forbidden(text: str) -> list:
    """מחזיר רשימת ממצאים: (line_no, col, char, char_name, context)."""
    findings = []
    lines = text.split("\n")
    for line_no, line in enumerate(lines, start=1):
        for col, ch in enumerate(line):
            if ch in FORBIDDEN_CHARS:
                start = max(0, col - 20)
                end = min(len(line), col + 21)
                context = line[start:end]
                findings.append(
                    {
                        "line": line_no,
                        "col": col + 1,
                        "char": ch,
                        "name": FORBIDDEN_CHARS[ch],
                        "context": context,
                    }
                )
    return findings


def main():
    # קלט: argv או stdin
    if len(sys.argv) > 1:
        path = Path(sys.argv[1])
        if not path.exists():
            print(f"ERROR: file not found: {path}", file=sys.stderr)
            sys.exit(2)
        text = read_text_from_path(path)
        source_label = str(path)
    else:
        if sys.stdin.isatty():
            print(__doc__, file=sys.stderr)
            sys.exit(2)
        text = sys.stdin.read()
        source_label = "<stdin>"

    findings = find_forbidden(text)

    if not findings:
        print(f"OK: no em-dash or similar forbidden characters in {source_label}")
        sys.exit(0)

    print(
        f"FAIL: found {len(findings)} forbidden character(s) in {source_label}",
        file=sys.stderr,
    )
    print("", file=sys.stderr)
    for f in findings:
        print(
            f"  line {f['line']}, col {f['col']}: {f['name']}",
            file=sys.stderr,
        )
        print(f"    context: ...{f['context']}...", file=sys.stderr)
        print("", file=sys.stderr)
    print(SUGGESTIONS, file=sys.stderr)
    sys.exit(1)


if __name__ == "__main__":
    main()
